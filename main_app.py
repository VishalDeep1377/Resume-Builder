"""
AI-Powered Resume Evaluator - Main Application
==============================================

This is the main Streamlit application that provides a user-friendly interface
for the AI-powered resume evaluator. It integrates all the modules and provides
a comprehensive analysis of resume quality, ATS-friendliness, and job matching.

Features:
- PDF resume upload and parsing
- ATS-friendliness analysis
- Job description matching
- Resume quality scoring
- Detailed feedback and suggestions
- Clean, modern UI

Author: AI Resume Evaluator
"""

import streamlit as st
import os
import tempfile
import logging
from typing import Dict, Any, Optional
import pandas as pd
import openai
import plotly.graph_objects as go
import plotly.express as px

# Import our custom modules
from resume_parser import ResumeParser
from ats_checker import ATSChecker
from jd_matcher import JobDescriptionMatcher
from scoring_model import ResumeScoringModel
from utils import (
    validate_pdf_file, format_score, get_score_color, create_summary_report,
    create_detailed_breakdown, validate_job_description, truncate_text
)
from chatbot_assistant import chat_with_resume_assistant

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Page configuration
st.set_page_config(
    page_title="AI Resume Evaluator",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Add mobile-friendly CSS at the top (after st.set_page_config)
st.markdown('''
<style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;600;700&display=swap');

    :root {
        --primary-color: #4A90E2;
        --primary-hover-color: #357ABD;
        --background-color: #0F172A;
        --secondary-background-color: #1E293B;
        --text-color: #E2E8F0;
        --header-color: #FFFFFF;
        --border-color: #334155;
    }

    html, body, [class*="st-"], [class*="css"] {
        font-family: 'Poppins', sans-serif;
        color: var(--text-color);
    }

    .stApp {
        background-color: var(--background-color);
    }

    .main-header {
        color: var(--header-color);
        text-align: center;
        font-weight: 700;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
    }

    [data-testid="stSidebar"] {
        background-color: var(--secondary-background-color);
        border-right: 1px solid var(--border-color);
    }

    .stButton>button {
        color: var(--header-color);
        background: linear-gradient(90deg, var(--primary-color), #3E7AC7);
        border: none;
        border-radius: 8px;
        padding: 12px 24px;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2);
    }

    .stButton>button:hover {
        background: linear-gradient(90deg, var(--primary-hover-color), #356CB0);
        box-shadow: 0 6px 20px rgba(0, 0, 0, 0.3);
        transform: translateY(-2px);
    }

    .stDownloadButton>button {
        color: var(--primary-color);
        background-color: transparent;
        border: 1px solid var(--primary-color);
        border-radius: 8px;
        padding: 12px 24px;
        font-weight: 600;
        transition: all 0.3s ease;
    }

    .stDownloadButton>button:hover {
        color: var(--header-color);
        background-color: var(--primary-color);
    }

    .stTextInput>div>input, .stTextArea>div>textarea {
        background-color: var(--secondary-background-color);
        color: var(--text-color);
        border: 1px solid var(--border-color);
        border-radius: 8px;
    }

    .stTabs [data-baseweb="tab-list"] button {
        background-color: transparent;
        color: var(--text-color);
        border-bottom: 2px solid transparent;
        transition: all 0.3s ease;
        font-size: 1.1rem;
    }

    .stTabs [data-baseweb="tab-list"] button[aria-selected="true"] {
        color: var(--primary-color);
        border-bottom: 2px solid var(--primary-color);
    }

    [data-testid="stMetric"] {
        background-color: var(--secondary-background-color);
        border-radius: 10px;
        padding: 1rem;
        border-left: 4px solid var(--primary-color);
    }

    [data-testid="stExpander"] {
        background-color: var(--secondary-background-color);
        border-radius: 8px;
        border: 1px solid var(--border-color);
    }

    [data-testid="stExpander"] summary {
        font-weight: 600;
        color: var(--primary-color);
    }
    
    .score-card {
        background-color: var(--secondary-background-color);
        border-radius: 10px;
        padding: 20px;
        text-align: center;
        border-left: 5px solid var(--primary-color);
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }

    .score-card h3 {
        color: var(--header-color);
        margin-bottom: 10px;
    }
    
</style>
''', unsafe_allow_html=True)

os.environ["HOME"] = "."
os.environ["STREAMLIT_HOME"] = "."

class ResumeEvaluatorApp:
    """
    Main application class for the AI Resume Evaluator.
    
    This class manages the Streamlit interface and coordinates all
    the analysis modules to provide a comprehensive resume evaluation.
    """
    
    def __init__(self):
        """Initialize the application and load models"""
        self.parser = None
        self.ats_checker = None
        self.jd_matcher = None
        self.scoring_model = None
        
        # Initialize session state
        if 'analysis_complete' not in st.session_state:
            st.session_state.analysis_complete = False
        if 'parsed_resume' not in st.session_state:
            st.session_state.parsed_resume = None
        if 'analysis_results' not in st.session_state:
            st.session_state.analysis_results = None
    
    def initialize_models(self):
        """Initialize all the analysis models"""
        try:
            with st.spinner("Loading AI models..."):
                # Initialize parser
                self.parser = ResumeParser()
                
                # Initialize ATS checker
                self.ats_checker = ATSChecker()
                
                # Initialize job description matcher
                self.jd_matcher = JobDescriptionMatcher()
                
                # Initialize and train scoring model
                self.scoring_model = ResumeScoringModel()
                
                # Check if pre-trained model exists
                model_path = "models/resume_model.pkl"
                if os.path.exists(model_path):
                    self.scoring_model.load_model(model_path)
                    st.success("Pre-trained model loaded successfully!")
                else:
                    with st.spinner("Training resume scoring model..."):
                        self.scoring_model.train_model()
                        # Create models directory if it doesn't exist
                        os.makedirs("models", exist_ok=True)
                        self.scoring_model.save_model(model_path)
                    st.success("Model trained and saved successfully!")
                
            return True
            
        except Exception as e:
            st.error(f"Error initializing models: {str(e)}")
            st.error("Please check that all required packages are installed.")
            return False
    
    def render_header(self):
        """Render the application header"""
        st.markdown('<h1 class="main-header">AI Resume Evaluator</h1>', unsafe_allow_html=True)
        st.markdown("""
        <div style="text-align: center; margin-bottom: 2rem;">
            <p style="font-size: 1.2rem; color: #666;">
                Advanced AI-powered resume analysis for ATS optimization and job matching
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    def render_sidebar(self):
        """Render the sidebar with information and instructions"""
        with st.sidebar:
            st.header("Quick Start")
            st.markdown("""
1. **Upload** your PDF resume
2. **Paste Job Description**
3. **Analyze** your resume
4. **See Results & Suggestions**

**PDF only** · **Max 10MB**
""")
            st.markdown("---")
            st.header("Features")
            st.markdown("""
- **AI Resume Check**: ATS, job match, quality score
- **Section Feedback**: Smart tips for every part
- **Job Match**: See missing keywords instantly
- **AI Rewrite**: Improve any section with GPT
- **Cover Letter**: 1-click, job-tailored letter
- **LinkedIn Check**: Compare resume & profile
- **Download Improved Resume**
- **Mobile Friendly**
""")
            st.markdown("---")
            st.header("AI Features Setup")
            st.markdown("""
To use AI rewriting, cover letter, or chatbot features:

1. **Create/Open an OpenAI Account**  
   Go to [OpenAI API Keys](https://platform.openai.com/account/api-keys)
2. **Sign In or Sign Up**  
   Use your email, Google, or Microsoft account
3. **Click 'Create new secret key'**  
   (Button is at the top right)
4. **Copy the key**  
   (It starts with `sk-...`)
5. **Paste it in the sidebar field below**
6. *(If you run out of credits, add a payment method in OpenAI billing)*

*Your key is only used on your device and never stored.*
""")
            
            st.header("OpenAI API Key (for AI Rewrite)")
            openai_api_key = st.text_input(
                "Enter your OpenAI API key",
                type="password",
                help="Get your key from https://platform.openai.com/account/api-keys"
            )
            st.session_state.openai_api_key = openai_api_key
    
    def upload_resume(self) -> Optional[str]:
        """Handle resume file upload"""
        st.header("Upload Your Resume")
        
        uploaded_file = st.file_uploader(
            "Choose a PDF file",
            type=['pdf'],
            help="Upload your resume in PDF format (max 10MB)"
        )
        
        if uploaded_file is not None:
            # Validate file type and size
            if uploaded_file.type != "application/pdf":
                st.error("File is not a PDF")
                return None
            elif uploaded_file.size > 10 * 1024 * 1024:
                st.error("File is too large (max 10MB)")
                return None
            else:
                # Save uploaded file to a temporary file for compatibility
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                    tmp_file.write(uploaded_file.read())
                    temp_path = tmp_file.name
                # Display file info
                file_size = uploaded_file.size / (1024 * 1024)  # MB
                st.success("File uploaded successfully!")
                st.info(f"File: {uploaded_file.name}")
                st.info(f"Size: {file_size:.2f} MB")
                return temp_path
        
        return None
    
    def input_job_descriptions(self) -> str:
        """Allow user to input or upload a single job description"""
        st.header("Job Description")
        jd_text = ""
        with st.expander("Job Description"):
            jd_text = st.text_area(
                "Paste job description",
                key="jd_text_single",
                height=150,
                placeholder="Paste the full job description here..."
            )
            jd_file = st.file_uploader(
                "Or upload JD (TXT or PDF)",
                type=["txt", "pdf"],
                key="jd_file_single"
            )
            if jd_file is not None:
                if jd_file.type == "text/plain":
                    jd_text = jd_file.read().decode("utf-8")
                elif jd_file.type == "application/pdf":
                    import pdfplumber
                    with pdfplumber.open(jd_file) as pdf:
                        jd_text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        return jd_text.strip()

    def analyze_resume(self, pdf_path: str, job_description: str):
        """Perform comprehensive resume analysis for a single JD"""
        try:
            progress_bar = st.progress(0)
            status_text = st.empty()
            status_text.text("Parsing resume...")
            progress_bar.progress(10)
            parsed_resume = self.parser.parse_resume(pdf_path)
            st.session_state.parsed_resume = parsed_resume
            status_text.text("Analyzing ATS-friendliness...")
            progress_bar.progress(25)
            ats_report = self.ats_checker.analyze_ats_friendliness(parsed_resume, job_description)
            progress_bar.progress(35)
            status_text.text("Matching with Job Description...")
            jd_report = self.jd_matcher.match_resume_to_job(parsed_resume['raw_text'], job_description)
            progress_bar.progress(75)
            status_text.text("Scoring resume quality...")
            features = self.scoring_model.extract_features(parsed_resume)
            scoring_result = self.scoring_model.predict_score(features)
            progress_bar.progress(100)
            status_text.text("Analysis complete!")
            st.session_state.analysis_results = {
                'ats_report': ats_report,
                'jd_report': jd_report,
                'scoring_result': scoring_result,
                'features': features,
                'job_description': job_description
            }
            st.session_state.analysis_complete = True
            try:
                os.unlink(pdf_path)
            except:
                pass
            st.success("Resume analysis completed successfully!")
        except Exception as e:
            st.error(f"Error during analysis: {str(e)}")
            logger.error(f"Analysis error: {str(e)}")

    def display_results(self):
        if not st.session_state.analysis_complete or not st.session_state.analysis_results:
            return
        results = st.session_state.analysis_results
        parsed_resume = st.session_state.parsed_resume
        summary = create_summary_report(
            results['ats_report'],
            results['jd_report'],
            results['scoring_result']
        )
        st.header("Overall Assessment")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric(
                "ATS Score",
                f"{results['ats_report'].overall_score:.1f}",
                delta=None,
                delta_color="normal"
            )
        with col2:
            st.metric(
                "Job Match",
                f"{results['jd_report'].overall_similarity * 100:.1f}%",
                delta=None,
                delta_color="normal"
            )
        with col3:
            st.metric(
                "Quality Score",
                f"{results['scoring_result'].overall_score:.1f}",
                delta=None,
                delta_color="normal"
            )
        with col4:
            avg_score = (results['ats_report'].overall_score + 
                        results['jd_report'].overall_similarity * 100 + 
                        results['scoring_result'].overall_score) / 3
            st.metric(
                "Average Score",
                f"{avg_score:.1f}",
                delta=None,
                delta_color="normal"
            )
        tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
            "ATS", "Match", "Quality", "Content", "Tips", "Cover", "Assistant"
        ])
        with tab1:
            self.display_ats_analysis(results['ats_report'])
        with tab2:
            self.display_jd_matching(results['jd_report'])
        with tab3:
            self.display_quality_score(results['scoring_result'])
        with tab4:
            self.display_content_analysis(parsed_resume, results['scoring_result'])
        with tab5:
            self.display_recommendations(summary)
        with tab6:
            self.display_cover_letter(parsed_resume, results['job_description'])
        with tab7:
            resume_text = parsed_resume.get('raw_text', '') if parsed_resume else ''
            chat_with_resume_assistant(resume_text)
    
    def display_ats_analysis(self, ats_report):
        """Display ATS-friendliness analysis"""
        st.subheader("ATS-Friendliness Analysis")
        
        # Overall score
        score_color = get_score_color(ats_report.overall_score)
        st.markdown(f"""
        <div class="score-card">
            <h3>Overall ATS Score: {ats_report.overall_score:.1f}/100</h3>
        </div>
        """, unsafe_allow_html=True)
        
        # Individual checks
        st.subheader("Detailed Checks")
        
        for check in ats_report.checks:
            with st.expander(f"{check.name} - {check.score:.1f}/100"):
                st.write(f"**Status:** {check.status.value}")
                st.write(f"**Description:** {check.description}")
                
                if check.suggestions:
                    st.write("**Suggestions:**")
                    for suggestion in check.suggestions:
                        st.write(f"• {suggestion}")
        
        # Summary statistics
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Passed Checks", ats_report.passed_checks)
        with col2:
            st.metric("Failed Checks", ats_report.failed_checks)
        with col3:
            st.metric("Warnings", ats_report.warnings)
    
    def display_quality_score(self, scoring_result):
        """Display resume quality scoring"""
        st.subheader("Resume Quality Score")
        
        # Overall score
        score_color = get_score_color(scoring_result.overall_score)
        st.markdown(f"""
        <div class="score-card">
            <h3>Overall Quality Score: {scoring_result.overall_score:.1f}/100</h3>
            <p>Confidence: {scoring_result.confidence:.1%}</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Feature breakdown
        st.subheader("Feature Breakdown")
        
        # Create a more visual breakdown with progress bars
        for feature, score in scoring_result.feature_scores.items():
            st.markdown(f"**{feature.replace('_', ' ').title()}**")
            st.progress(int(score) / 100)

        # Detailed breakdown
        st.subheader("Detailed Analysis")
        for category, description in scoring_result.breakdown.items():
            st.write(f"**{category}:** {description}")
    
    def display_content_analysis(self, parsed_resume, scoring_result):
        """Display content analysis with section-by-section feedback"""
        st.subheader("Content Analysis & Section Feedback")

        # Helper: Find section by title
        def get_section_content(sections, key):
            for section in sections:
                if section.title.lower() == key:
                    return section.content
            return None

        sections = parsed_resume.get('sections', [])
        found_sections = {s.title.lower() for s in sections}
        required_sections = ['skills', 'experience', 'education']
        section_labels = {
            'skills': 'Skills',
            'experience': 'Work Experience',
            'education': 'Education',
        }
        section_tips = {
            'skills': "Include a mix of technical and soft skills relevant to the job. Use keywords from the job description.",
            'experience': "List your roles in reverse-chronological order. Use bullet points and quantify achievements.",
            'education': "Include your degree, university, and graduation year. Add honors or relevant coursework if possible.",
        }

        # Section-by-section feedback
        for key in required_sections:
            label = section_labels[key]
            st.markdown(f"### {label}")
            content = get_section_content(sections, key)
            if content:
                with st.expander(f"View {label} Section"):
                    st.write(content)
                    st.info(section_tips[key])
                    # Inline feedback (example: check for length, keywords, etc.)
                    if key == 'skills':
                        skills = parsed_resume.get('skills', [])
                        if len(skills) < 5:
                            st.warning("Consider adding more skills to strengthen your profile.")
                        else:
                            st.success(f"{len(skills)} skills detected.")
                    elif key == 'experience':
                        experience = parsed_resume.get('experience', [])
                        if not experience:
                            st.warning("No detailed experience entries found. Add more details for each job.")
                        else:
                            st.success(f"{len(experience)} experience entries detected.")
                    elif key == 'education':
                        st.info("Education section found. Make sure to include your most recent degree.")
            else:
                st.error(f"{label} section not found! Add this section for a stronger resume.")
                st.info(section_tips[key])

        # Show other sections if present
        extra_sections = [s for s in sections if s.title.lower() not in required_sections]
        if extra_sections:
            st.markdown("### Additional Sections")
            for section in extra_sections:
                with st.expander(f"{section.title.capitalize()}"):
                    st.write(section.content)

        # Statistics (unchanged)
        if 'raw_text' in parsed_resume:
            st.write("**Document Statistics:**")
            text = parsed_resume['raw_text']
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Characters", len(text))
                st.metric("Words", len(text.split()))
            with col2:
                st.metric("Sentences", len(text.split('.')))
                st.metric("Paragraphs", len([p for p in text.split('\n\n') if p.strip()]))
            with col3:
                st.metric("Skills Found", len(parsed_resume.get('skills', [])))
                st.metric("Experience Entries", len(parsed_resume.get('experience', [])))
        
        # Feature importance
        st.subheader("Feature Importance")
        st.markdown("This chart shows which resume features had the biggest impact on your quality score.")
        if scoring_result.feature_importance:
            importance_df = pd.DataFrame({
                'Feature': list(scoring_result.feature_importance.keys()),
                'Importance': list(scoring_result.feature_importance.values())
            }).sort_values('Importance', ascending=False)
            
            st.bar_chart(importance_df.set_index('Feature'))
        else:
            st.info("Feature importance data is not available.")
    
    def display_jd_matching(self, jd_report):
        """Display job description matching analysis"""
        st.subheader("Job Description Match Analysis")

        st.markdown(f"""
        <div class="score-card">
            <h3>Overall Match Score: {jd_report.overall_similarity * 100:.1f}%</h3>
            <p>This score reflects how well the keywords and phrases in your resume align with the job description.</p>
        </div>
        """, unsafe_allow_html=True)

        st.subheader("Keyword Breakdown")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Keywords in JD", jd_report.total_keywords)
        with col2:
            st.metric("Keywords Found", len(jd_report.matching_keywords))
        with col3:
            st.metric("Keywords Missing", len(jd_report.missing_keywords))

        tab1, tab2 = st.tabs(["Missing Keywords", "Matching Keywords"])

        with tab1:
            if jd_report.missing_keywords:
                st.error("Consider adding these keywords to your resume to improve your match score:")
                st.text(", ".join(jd_report.missing_keywords))
            else:
                st.success("Excellent! No important keywords are missing from your resume.")

        with tab2:
            if jd_report.matching_keywords:
                st.info("These keywords from the job description were found in your resume:")
                st.text(", ".join(jd_report.matching_keywords))
            else:
                st.warning("No matching keywords were found.")

        if jd_report.suggestions:
            st.subheader("Suggestions for Improvement")
            for suggestion in jd_report.suggestions:
                st.write(f"👉 {suggestion}")

        with st.expander("View Detailed Text Analysis"):
            st.text(jd_report.detailed_analysis)
    
    def display_recommendations(self, summary):
        """Display recommendations and suggestions"""
        st.subheader("Recommendations & Suggestions")
        
        # Priority actions
        if summary['priority_actions']:
            st.write("**Priority Actions:**")
            for action in summary['priority_actions']:
                st.write(f"• {action}")
        
        # All recommendations
        if summary['recommendations']:
            st.write("**All Recommendations:**")
            for rec in summary['recommendations']:
                st.write(f"• {rec}")
        
        # Download report
        st.subheader("Download Report")
        
        # Create a simple text report
        report_text = f"""
AI Resume Evaluator Report
Generated: {summary['timestamp']}

OVERALL ASSESSMENT
Status: {summary['overall_assessment']['status']}
Message: {summary['overall_assessment']['message']}

SCORES
{chr(10).join([f"{k}: {v}" for k, v in summary['scores'].items()])}

RECOMMENDATIONS
{chr(10).join([f"• {rec}" for rec in summary['recommendations']])}
        """
        
        st.download_button(
            label="Download Report",
            data=report_text,
            file_name="resume_evaluation_report.txt",
            mime="text/plain"
        )
        
        # Generate improved resume DOCX
        st.subheader("Download Improved Resume (DOCX)")
        if st.button("Generate Improved Resume (DOCX)", type="secondary"):
            with st.spinner("Generating improved resume..."):
                try:
                    improved_docx = self.generate_improved_resume_docx(st.session_state.parsed_resume, st.session_state.analysis_results['job_description'])
                    st.download_button(
                        label="Download Improved Resume (DOCX)",
                        data=improved_docx,
                        file_name="improved_resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("Improved resume generated successfully!")
                except Exception as e:
                    st.error(f"Error generating improved resume: {str(e)}")
    
    def generate_improved_resume_docx(self, parsed_resume, job_description):
        from docx import Document
        from docx.shared import Pt
        import openai
        openai_api_key = st.session_state.get('openai_api_key', None)
        doc = Document()
        # Header
        contact = parsed_resume.get('contact_info', {})
        name = getattr(contact, 'name', None) or "[Your Name]"
        email = getattr(contact, 'email', None) or "[Email]"
        phone = getattr(contact, 'phone', None) or "[Phone]"
        location = getattr(contact, 'location', None) or "[Location]"
        doc.add_heading(name, 0)
        doc.add_paragraph(f"Email: {email} | Phone: {phone} | Location: {location}")
        # Section helper
        def add_section(title, content):
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)
        # AI rewrite helper
        def ai_rewrite(section_title, section_text):
            if not openai_api_key or not section_text.strip():
                return section_text
            try:
                client = openai.OpenAI(api_key=openai_api_key)
                prompt = f"Rewrite this resume section to be more impactful, clear, and ATS-friendly for the following job description.\nJob Description:\n{job_description}\nSection:\n{section_text}"
                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a professional resume writer and career coach."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=400,
                    temperature=0.7,
                )
                return response.choices[0].message.content.strip()
            except Exception:
                return section_text
        # Sections
        for section in parsed_resume.get('sections', []):
            title = section.title.capitalize()
            content = section.content
            improved_content = ai_rewrite(title, content)
            add_section(title, improved_content)
        # Skills
        skills = parsed_resume.get('skills', [])
        if skills:
            add_section("Skills", ', '.join(skills))
        # Experience
        experience = parsed_resume.get('experience', [])
        if experience:
            doc.add_heading("Experience", level=1)
            for exp in experience:
                p = doc.add_paragraph()
                p.add_run(exp.get('title', 'Job Title')).bold = True
                p.add_run(f" | {exp.get('description', '')}")
        # Save to bytes
        from io import BytesIO
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        return f.getvalue()
    
    def input_linkedin_profile(self):
        """Let user upload LinkedIn PDF or paste profile text"""
        st.header("LinkedIn Profile Analyzer (Beta)")
        st.markdown("""
**How to use:**
- Export your LinkedIn profile as a PDF (Profile > More > Save to PDF) and upload it below, **or**
- Copy and paste your profile text from LinkedIn into the text area.

**Note:** Direct LinkedIn URLs are not supported. Please use PDF or text only.
""")
        linkedin_text = ""
        uploaded_file = st.file_uploader("Upload your LinkedIn PDF (exported from LinkedIn)", type=["pdf"], key="linkedin_pdf")
        if uploaded_file is not None:
            import pdfplumber
            with pdfplumber.open(uploaded_file) as pdf:
                linkedin_text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        manual_text = st.text_area("Or paste your LinkedIn profile text here", key="linkedin_text", height=200)
        if manual_text.strip():
            linkedin_text = manual_text.strip()
        return linkedin_text

    def compare_linkedin_resume(self, resume_data, linkedin_text):
        """Compare parsed resume with LinkedIn profile text and return mismatches and a consistency score"""
        import re
        # Extract skills from LinkedIn
        skills_section = re.search(r'Skills[\s\S]+?(?=\n\w|$)', linkedin_text, re.I)
        linkedin_skills = []
        if skills_section:
            linkedin_skills = re.findall(r'\b[A-Za-z][A-Za-z0-9\- ]{2,}\b', skills_section.group())
        # Extract positions from LinkedIn
        positions = re.findall(r'(?:Experience|Positions)[\s\S]+?(?=\n\w|$)', linkedin_text, re.I)
        linkedin_positions = []
        for pos in positions:
            linkedin_positions += re.findall(r'\b[A-Z][A-Za-z0-9\- ]{2,}\b', pos)
        # Extract summary
        summary_match = re.search(r'(Summary|About)[\s\S]+?(?=\n\w|$)', linkedin_text, re.I)
        linkedin_summary = summary_match.group().strip() if summary_match else ""
        # Resume data
        resume_skills = set(resume_data.get('skills', []))
        resume_experience = resume_data.get('experience', [])
        resume_positions = set([e.get('title', '').strip() for e in resume_experience if e.get('title')])
        resume_summary = ""
        for s in resume_data.get('sections', []):
            if s.title.lower() in ['summary', 'profile', 'objective']:
                resume_summary = s.content.strip()
                break
        # Compare
        missing_in_resume = set(linkedin_skills) - resume_skills
        missing_in_linkedin = resume_skills - set(linkedin_skills)
        position_mismatches = resume_positions.symmetric_difference(set(linkedin_positions))
        summary_match = 1 if resume_summary and linkedin_summary and resume_summary[:50].lower() in linkedin_summary.lower() else 0
        # Consistency score
        skill_score = 1 - (len(missing_in_resume) + len(missing_in_linkedin)) / (len(resume_skills | set(linkedin_skills)) + 1e-5)
        position_score = 1 - (len(position_mismatches) / (len(resume_positions | set(linkedin_positions)) + 1e-5))
        consistency_score = (0.5 * skill_score + 0.3 * position_score + 0.2 * summary_match)
        return {
            'missing_in_resume': list(missing_in_resume),
            'missing_in_linkedin': list(missing_in_linkedin),
            'position_mismatches': list(position_mismatches),
            'summary_match': summary_match,
            'consistency_score': consistency_score
        }

    def display_linkedin_comparison(self, comparison):
        st.subheader("LinkedIn vs Resume Consistency")
        st.metric("Consistency Score", f"{comparison['consistency_score']*100:.1f}%")
        st.markdown(f"**Summary Match:** {'Yes' if comparison['summary_match'] else 'No'}")
        st.markdown(f"**Skills in LinkedIn but not Resume:** {', '.join(comparison['missing_in_resume']) if comparison['missing_in_resume'] else 'None!'}")
        st.markdown(f"**Skills in Resume but not LinkedIn:** {', '.join(comparison['missing_in_linkedin']) if comparison['missing_in_linkedin'] else 'None!'}")
        st.markdown(f"**Position Mismatches:** {', '.join(comparison['position_mismatches']) if comparison['position_mismatches'] else 'None!'}")

    def display_cover_letter(self, parsed_resume, job_description):
        st.subheader("Cover Letter Generator")
        openai_api_key = st.session_state.get('openai_api_key', None)
        # Extract info for prompt
        name = parsed_resume.get('contact_info', {}).name or "[Your Name]"
        skills = ', '.join(parsed_resume.get('skills', [])[:8])
        experience = parsed_resume.get('experience', [])
        exp_summary = experience[0]['description'] if experience else "relevant experience"
        job_title = job_description.split('\n')[0][:60] if job_description else "the position"
        # Prompt for GPT
        prompt = f"""
Write a professional, tailored cover letter for the following job:

Job Description:
{job_description}

Resume Summary:
Name: {name}
Skills: {skills}
Experience: {exp_summary}

The letter should:
- Address the hiring manager
- Highlight relevant skills and experience
- Show enthusiasm for the role
- Be concise (max 300 words)
"""
        cover_letter = ""
        if openai_api_key:
            import openai
            try:
                client = openai.OpenAI(api_key=openai_api_key)
                with st.spinner("Generating cover letter with GPT..."):
                    response = client.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": "You are a professional career coach and cover letter writer."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=400,
                        temperature=0.7,
                    )
                    cover_letter = response.choices[0].message.content.strip()
            except Exception as e:
                st.error(f"OpenAI API error: {e}")
        if not cover_letter:
            # Fallback template
            cover_letter = f"""
Dear Hiring Manager,

I am excited to apply for {job_title}. With skills in {skills}, and experience in {exp_summary}, I am confident I can contribute to your team. I am passionate about this opportunity and eager to bring my expertise to your organization.

Thank you for considering my application.

Sincerely,
{name}
"""
        st.text_area("Generated Cover Letter", cover_letter, height=300)
        st.download_button(
            label="Download Cover Letter",
            data=cover_letter,
            file_name="cover_letter.txt",
            mime="text/plain"
        )
        st.info("You can copy or download your tailored cover letter above.")

    def run(self):
        """Main application runner"""
        # Initialize models
        if not self.initialize_models():
            st.error("Failed to initialize models. Please check your installation.")
            return
        
        # Render header and sidebar
        self.render_header()
        self.render_sidebar()
        
        # Upload section
        pdf_path = self.upload_resume()

        # Job description section
        job_description = self.input_job_descriptions()

        # ATS-Friendly Resume Templates Section
        st.header("Download ATS-Friendly Resume Templates")
        st.markdown("""
        Choose from professionally designed, ATS-safe resume templates. These templates use standard fonts, clear section headings, and simple layouts to maximize compatibility with Applicant Tracking Systems (ATS).
        """)
        tab1, tab2, tab3 = st.tabs(["Minimal Template", "Modern Template", "Student Template"])
        with tab1:
            with open("data/sample_resumes/ATS_Minimal_Template.docx", "rb") as f:
                st.download_button(
                    label="Download DOCX",
                    data=f.read(),
                    file_name="ATS_Minimal_Template.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            st.markdown("**Minimal Template**: Clean, simple design perfect for ATS systems. Includes standard sections with clear formatting.")
        with tab2:
            with open("data/sample_resumes/ATS_Modern_Template.docx", "rb") as f:
                st.download_button(
                    label="Download DOCX",
                    data=f.read(),
                    file_name="ATS_Modern_Template.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            st.markdown("**Modern Template**: Professional layout with enhanced sections for experienced professionals.")
        with tab3:
            with open("data/sample_resumes/ATS_Student_Template.docx", "rb") as f:
                st.download_button(
                    label="Download DOCX",
                    data=f.read(),
                    file_name="ATS_Student_Template.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            st.markdown("**Student Template**: Designed for students and recent graduates with focus on projects and academic achievements.")
        st.markdown("---")
        
        # Main content area
        if not st.session_state.analysis_complete:
            # Analysis button
            if pdf_path and job_description:
                st.header("Start Analysis")
                
                if st.button("Analyze Resume", type="primary", use_container_width=True):
                    self.analyze_resume(pdf_path, job_description)
                    st.rerun()
            else:
                st.info("Please upload a resume and enter a job description to begin the analysis.")
        
        else:
            # Display results
            self.display_results()
            
            # Reset button
            if st.button("Start New Analysis", use_container_width=True):
                st.session_state.analysis_complete = False
                st.session_state.parsed_resume = None
                st.session_state.analysis_results = None
                st.rerun()

# Main application entry point
if __name__ == "__main__":
    app = ResumeEvaluatorApp()
    app.run() 